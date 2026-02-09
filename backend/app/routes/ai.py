"""
Proxy endpoints for AI Interviewer microservice integration.
These endpoints forward requests to the AI service (Node.js/Genkit) and return responses to the frontend.
"""

import os
import logging
import httpx
from fastapi import APIRouter, Request, HTTPException, status, Depends
from fastapi.responses import JSONResponse
from sqlalchemy import select
from app.models.ai_report import AIReport
from app.core.config import settings
from app.middleware.auth import get_current_user
from app.core.database import get_db
from sqlalchemy.ext.asyncio import AsyncSession
from app.services.ai_report_service import AIReportService
import httpx
import base64
from fastapi import UploadFile, File, Form

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/v1/ai", tags=["AI Interviewer"])

# AI service base URL (internal Docker network or env var)
# The AI service runs as Genkit embedded in the Next.js app or as a separate service
# For now, we point to the Next.js app which has Genkit API routes
AI_SERVICE_URL = os.getenv("AI_SERVICE_URL", "http://ai-service:3000/api")

async def proxy_to_ai_service(path: str, method: str = "POST", data=None, params=None, headers=None):
    url = f"{AI_SERVICE_URL}{path}"
    async with httpx.AsyncClient(timeout=30) as client:
        try:
            response = await client.request(
                method,
                url,
                json=data,
                params=params,
                headers=headers,
            )
            response.raise_for_status()
            return JSONResponse(status_code=response.status_code, content=response.json())
        except httpx.HTTPStatusError as e:
            raise HTTPException(status_code=e.response.status_code, detail=e.response.text)
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

# Example: Proxy for AI Interview Session Generation
@router.post("/interview-session")
async def create_interview_session(request: Request, user=Depends(get_current_user)):
    body = await request.json()
    return await proxy_to_ai_service("/interview/sessions", method="POST", data=body)

# Example: Proxy for Resume Analysis
@router.post("/resume-analysis")
async def resume_analysis(request: Request, user=Depends(get_current_user)):
    body = await request.json()
    return await proxy_to_ai_service("/interview/{token}/get-resume-analysis", method="POST", data=body)


@router.post("/ats")
async def ats_for_interview(
    request: Request,
    session: AsyncSession = Depends(get_db),
):
    """ATS check for interview flow - saves score to interview record and creates AI report.
    
    If candidate already has an ATS score (from dashboard), returns that instead of running again.
    
    Expected body: { resume_text, job_description (optional), interview_id (optional) }
    Returns: { score, summary, highlights, improvements, keywords_found, keywords_missing, verdict, from_cache }
    """
    try:
        import json as json_lib
        payload = await request.json()
        resume_text = payload.get("resume_text", "")
        job_description = payload.get("job_description", "")
        interview_id = payload.get("interview_id")
        
        # Check if interview already has ATS score (from dashboard or previous check)
        if interview_id:
            try:
                from app.models.candidate import Interview, Candidate
                interview = await session.get(Interview, interview_id)
                if interview:
                    # First check if interview already has ATS score
                    if interview.ats_score is not None and interview.ats_score > 0:
                        logger.info(f"Using cached ATS score {interview.ats_score} from interview {interview_id}")
                        score = interview.ats_score
                        if score >= 80:
                            verdict = "EXCELLENT"
                        elif score >= 70:
                            verdict = "GOOD"
                        elif score >= 60:
                            verdict = "FAIR"
                        else:
                            verdict = "NEEDS_IMPROVEMENT"
                        return {
                            "score": score,
                            "summary": "Resume analysis was completed earlier. Using saved score.",
                            "verdict": verdict,
                            "highlights": [],
                            "improvements": [],
                            "keywords_found": [],
                            "keywords_missing": [],
                            "from_cache": True
                        }
                    
                    # Check if candidate has ATS score from dashboard
                    candidate = await session.get(Candidate, interview.candidate_id)
                    if candidate and candidate.ats_score is not None and candidate.ats_score > 0:
                        logger.info(f"Using ATS score {candidate.ats_score} from candidate dashboard check")
                        # Copy to interview for consistency
                        interview.ats_score = candidate.ats_score
                        await session.commit()
                        
                        # Try to parse stored report
                        cached_report = {}
                        if candidate.ats_report:
                            try:
                                cached_report = json_lib.loads(candidate.ats_report)
                            except:
                                pass
                        
                        return {
                            "score": candidate.ats_score,
                            "summary": cached_report.get("summary", "Resume analysis was completed from dashboard. Using saved score."),
                            "verdict": cached_report.get("verdict", "FAIR"),
                            "highlights": cached_report.get("highlights", []),
                            "improvements": cached_report.get("improvements", []),
                            "keywords_found": cached_report.get("keywords_found", []),
                            "keywords_missing": cached_report.get("keywords_missing", []),
                            "from_cache": True
                        }
            except Exception as e:
                logger.warning(f"Error checking cached ATS: {e}")
        
        if not resume_text:
            raise HTTPException(status_code=400, detail="Missing resume_text")
        
        # Use AI service to generate enhanced ATS report
        from app.services.ai_service import generate_ats_report_enhanced
        
        result = await generate_ats_report_enhanced(resume_text, job_description)
        
        # Add verdict based on score
        score = result.get("score", 0)
        if score >= 80:
            result["verdict"] = "EXCELLENT"
        elif score >= 70:
            result["verdict"] = "GOOD"
        elif score >= 60:
            result["verdict"] = "FAIR"
        else:
            result["verdict"] = "NEEDS_IMPROVEMENT"
        
        result["from_cache"] = False
        
        # Save to interview if interview_id provided
        if interview_id:
            try:
                from app.models.candidate import Interview, Candidate
                interview = await session.get(Interview, interview_id)
                if interview:
                    interview.ats_score = score
                    interview.resume_text = resume_text[:10000]  # Store first 10k chars
                    
                    # Also save to candidate record for future interviews
                    candidate = await session.get(Candidate, interview.candidate_id)
                    if candidate:
                        candidate.ats_score = score
                        candidate.ats_report = json_lib.dumps(result)
                    
                    await session.commit()
                    
                    # Also create AI report for tracking
                    report = await AIReportService.create_report(
                        session=session,
                        company_id=interview.company_id,
                        report_type="ats_check",
                        provider_response=result,
                        candidate_id=interview.candidate_id,
                        interview_id=interview_id,
                        score=score,
                        summary=result.get("summary"),
                    )
                    await session.commit()
                    result["report_id"] = str(report.id)
            except Exception as e:
                logger.warning(f"Failed to save ATS score to interview: {e}")
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"ATS check failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"ATS check failed: {str(e)}")

@router.post("/ats-check")
async def ats_check(
    resume: UploadFile = File(None),
    job_description: str = Form(None),
    candidate_id: str = Form(None),
    user=Depends(get_current_user),
    session: AsyncSession = Depends(get_db),
):
    """ATS check for resume files or text - for quick candidate self-check.
    
    Accepts file upload (resume) or JSON body with resume_text.
    If user is a candidate, automatically saves the ATS score to their candidate record.
    Returns: { score, summary, highlights, improvements, keywords_found, keywords_missing, verdict }
    """
    try:
        import json
        from app.models.user import UserRole
        from app.models.candidate import Candidate
        from sqlalchemy import select
        
        resume_text = ""
        job_desc = job_description or ""
        
        # Auto-detect candidate_id from user email if user is a candidate
        actual_candidate_id = candidate_id
        if not actual_candidate_id and user and user.role == UserRole.CANDIDATE:
            # Look up candidate by email
            candidate_query = select(Candidate).filter(Candidate.email == user.email)
            result = await session.execute(candidate_query)
            candidate = result.scalars().first()
            if candidate:
                actual_candidate_id = str(candidate.id)
                logger.info(f"Auto-detected candidate_id {actual_candidate_id} for user {user.email}")
        
        # Handle file upload
        if resume:
            content = await resume.read()
            filename = resume.filename.lower() if resume.filename else ""
            
            # Handle text files
            if filename.endswith('.txt') or resume.content_type == 'text/plain':
                resume_text = content.decode('utf-8', errors='ignore')
            
            # Handle PDF files
            elif filename.endswith('.pdf') or resume.content_type == 'application/pdf':
                try:
                    import PyPDF2
                    import io
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                    for page in pdf_reader.pages:
                        resume_text += page.extract_text() or ""
                except Exception as pdf_err:
                    logger.warning(f"PyPDF2 failed: {pdf_err}, trying pdfplumber")
                    try:
                        import pdfplumber
                        import io
                        with pdfplumber.open(io.BytesIO(content)) as pdf:
                            for page in pdf.pages:
                                resume_text += page.extract_text() or ""
                    except Exception as plumber_err:
                        logger.error(f"pdfplumber also failed: {plumber_err}")
                        raise HTTPException(status_code=400, detail="Could not extract text from PDF. Please try a different file or paste the text directly.")
            
            # Handle Word documents
            elif filename.endswith('.docx') or resume.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                try:
                    import docx
                    import io
                    doc = docx.Document(io.BytesIO(content))
                    resume_text = "\n".join([para.text for para in doc.paragraphs])
                except Exception as docx_err:
                    logger.error(f"Failed to parse docx: {docx_err}")
                    raise HTTPException(status_code=400, detail="Could not extract text from Word document.")
            
            elif filename.endswith('.doc'):
                raise HTTPException(status_code=400, detail="Legacy .doc files are not supported. Please convert to .docx or PDF.")
            
            else:
                # Try to decode as text
                try:
                    resume_text = content.decode('utf-8', errors='ignore')
                except:
                    raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF, DOCX, or TXT.")
        
        if not resume_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from the resume. Please try a different file.")
        
        # Use AI service to generate ATS report
        from app.services.ai_service import generate_ats_report_enhanced
        
        result = await generate_ats_report_enhanced(resume_text.strip(), job_desc)
        
        # Add verdict based on score
        score = result.get("score", 0)
        if score >= 80:
            result["verdict"] = "EXCELLENT"
        elif score >= 70:
            result["verdict"] = "GOOD"
        elif score >= 60:
            result["verdict"] = "FAIR"
        else:
            result["verdict"] = "NEEDS_IMPROVEMENT"
        
        # Save to candidate record if actual_candidate_id is available (explicit or auto-detected)
        if actual_candidate_id:
            try:
                from uuid import UUID as UUID_TYPE
                cand_uuid = UUID_TYPE(actual_candidate_id)
                candidate = await session.get(Candidate, cand_uuid)
                if candidate:
                    candidate.ats_score = score
                    candidate.ats_report = json.dumps(result)
                    candidate.resume_text = resume_text.strip()[:50000]  # Save resume text (limit to 50KB)
                    await session.commit()
                    logger.info(f"Saved ATS score {score} to candidate {actual_candidate_id}")
            except Exception as e:
                logger.warning(f"Failed to save ATS score to candidate: {e}")
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"ATS check failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"ATS check failed: {str(e)}")


@router.post("/parse-resume")
async def parse_resume(
    file: UploadFile = File(...),
    user=Depends(get_current_user),
):
    """Parse a resume file (PDF, Word, TXT) and extract text content.
    
    Returns: { text: string }
    """
    try:
        content = await file.read()
        filename = file.filename.lower() if file.filename else ""
        
        # Handle text files
        if filename.endswith('.txt') or file.content_type == 'text/plain':
            text = content.decode('utf-8', errors='ignore')
            return {"text": text}
        
        # Handle PDF files
        if filename.endswith('.pdf') or file.content_type == 'application/pdf':
            try:
                import PyPDF2
                import io
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() or ""
                return {"text": text.strip()}
            except ImportError:
                # Fallback: try pdfplumber
                try:
                    import pdfplumber
                    import io
                    with pdfplumber.open(io.BytesIO(content)) as pdf:
                        text = ""
                        for page in pdf.pages:
                            text += page.extract_text() or ""
                    return {"text": text.strip()}
                except ImportError:
                    raise HTTPException(
                        status_code=500, 
                        detail="PDF parsing not available. Please paste your resume text instead."
                    )
        
        # Handle Word files
        if filename.endswith('.docx') or file.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            try:
                import docx
                import io
                doc = docx.Document(io.BytesIO(content))
                text = "\n".join([para.text for para in doc.paragraphs])
                return {"text": text.strip()}
            except ImportError:
                raise HTTPException(
                    status_code=500,
                    detail="Word document parsing not available. Please paste your resume text instead."
                )
        
        if filename.endswith('.doc'):
            raise HTTPException(
                status_code=400,
                detail="Legacy .doc format not supported. Please convert to .docx or paste the text."
            )
        
        raise HTTPException(
            status_code=400,
            detail="Unsupported file format. Please upload PDF, Word (.docx), or text file."
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Resume parsing failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to parse resume: {str(e)}")


# Example: Proxy for ATS Checker
@router.post("/ats-checker")
async def ats_checker(request: Request, user=Depends(get_current_user)):
    body = await request.json()
    return await proxy_to_ai_service("/ats/check", method="POST", data=body)


@router.post("/ats-checker-file")
async def ats_checker_file(
    file: UploadFile = File(...),
    candidate_email: str | None = None,
    candidate_id: str | None = None,
    user=Depends(get_current_user),
):
    """Accept a resume upload (PDF/DOCX/TXT) and forward to AI service for ATS checking.

    Because some AI services expect JSON, we encode the file as base64 and send metadata.
    The AI service is expected to accept a JSON body: { file_name, file_b64, candidate_email, candidate_id }
    If your AI service accepts multipart/form-data, this can be adapted later.
    """
    try:
        content = await file.read()
        file_b64 = base64.b64encode(content).decode('ascii')

        payload = {
            "file_name": file.filename,
            "file_b64": file_b64,
            "candidate_email": candidate_email,
            "candidate_id": candidate_id,
        }

        # Forward to AI service
        return await proxy_to_ai_service("/ats/check", method="POST", data=payload)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"ATS checker failed: {str(e)}")



@router.post("/ats-check-and-save")
async def ats_check_and_save(
    request: Request,
    session: AsyncSession = Depends(get_db),
    user=Depends(get_current_user),
):
    """Forward ATS check to AI service, persist the report into ai_reports and return saved record.

    Expected body: { candidate_id (optional), candidate_email (optional), file_b64 or resume_url or other fields }
    """
    try:
        payload = await request.json()

        # Use internal AIService adapter to call the provider synchronously
        from app.services.ai_service import generate_ats_report

        candidate_id = payload.get("candidate_id")
        company_id = user.company_id
        report_type = "ats"

        ai_report = await generate_ats_report(payload.get("resume_text") or payload.get("file_b64") or "")

        score = ai_report.get("score")
        summary = ai_report.get("summary")

        report = await AIReportService.create_report(
            session=session,
            company_id=company_id,
            report_type=report_type,
            provider_response=ai_report.get("raw") or ai_report,
            candidate_id=candidate_id,
            score=score,
            summary=summary,
            created_by=user.id,
        )
        await session.commit()

        return {"report_id": str(report.id), "ai_response": ai_report}

    except httpx.HTTPStatusError as e:
        raise HTTPException(status_code=e.response.status_code, detail=e.response.text)
    except Exception as e:
        await session.rollback()
        raise HTTPException(status_code=500, detail=f"ATS check and save failed: {str(e)}")


@router.get("/reports")
async def list_ai_reports(
    limit: int = 20,
    offset: int = 0,
    session: AsyncSession = Depends(get_db),
    user=Depends(get_current_user),
):
    """List AI reports for the current user's company.

    Returns a simple JSON structure: { reports: [ ... ] }
    """
    try:
        from sqlalchemy import or_
        
        # Include reports for the user's company OR reports with no company (general reports)
        stmt = (
            select(AIReport)
            .where(
                or_(
                    AIReport.company_id == user.company_id,
                    AIReport.company_id.is_(None)
                )
            )
            .order_by(AIReport.created_at.desc())
            .limit(limit)
            .offset(offset)
        )
        result = await session.execute(stmt)
        reports = result.scalars().all()

        # Simple serialization
        out = []
        for r in reports:
            out.append({
                "id": str(r.id),
                "report_type": r.report_type,
                "score": r.score,
                "summary": r.summary,
                "created_at": r.created_at.isoformat() if r.created_at else None,
            })

        return {"reports": out}
    except Exception as e:
        logger.error(f"Failed to list AI reports: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to list AI reports: {str(e)}")

# Example: Proxy for Transcript Fetching
@router.get("/transcript/{session_id}")
async def get_transcript(session_id: str, user=Depends(get_current_user)):
    return await proxy_to_ai_service(f"/reports/transcripts/{session_id}", method="GET")


@router.post("/transcript-callback")
async def transcript_callback(
    request: Request,
    session: AsyncSession = Depends(get_db),
    user=Depends(get_current_user),
):
    """Endpoint for AI service (or AssemblyAI proxy) to POST final transcript.

    Expected JSON body: { session_id, candidate_id, interview_id, transcript_text }
    This will create an ai_report of type 'transcript_verdict' by calling the ATS/generative adapter.
    """
    try:
        payload = await request.json()
        transcript = payload.get("transcript_text") or payload.get("transcript")
        candidate_id = payload.get("candidate_id")
        interview_id = payload.get("interview_id")

        if not transcript:
            raise HTTPException(status_code=400, detail="Missing transcript_text in payload")

        # Use AI adapter to generate verdict based on transcript
        from app.services.ai_service import generate_ats_report

        ai_report = await generate_ats_report(transcript)

        report = await AIReportService.create_report(
            session=session,
            company_id=user.company_id,
            report_type="transcript_verdict",
            provider_response=ai_report.get("raw") or ai_report,
            candidate_id=candidate_id,
            interview_id=interview_id,
            score=ai_report.get("score"),
            summary=ai_report.get("summary"),
            created_by=user.id,
        )
        await session.commit()

        return {"report_id": str(report.id), "ai_response": ai_report}

    except HTTPException:
        raise
    except Exception as e:
        await session.rollback()
        raise HTTPException(status_code=500, detail=f"Transcript callback failed: {str(e)}")


# ==================== COMPANY AI CONFIG ====================

from pydantic import BaseModel
from typing import Optional

class AIConfigUpdate(BaseModel):
    min_passing_score: Optional[int] = None
    min_ats_score: Optional[int] = None
    auto_reject_below: Optional[int] = None
    require_employee_review: Optional[bool] = None
    ats_enabled: Optional[bool] = None
    ai_verdict_enabled: Optional[bool] = None


@router.get("/config")
async def get_ai_config(
    session: AsyncSession = Depends(get_db),
    user=Depends(get_current_user),
):
    """
    Get company AI configuration settings.
    """
    try:
        from app.models.company_ai_config import CompanyAIConfig
        
        result = await session.execute(
            select(CompanyAIConfig).filter(CompanyAIConfig.company_id == user.company_id)
        )
        config = result.scalar_one_or_none()
        
        if config:
            return config.to_dict()
        else:
            # Return default config
            return {
                "id": None,
                "company_id": str(user.company_id),
                "min_passing_score": 70,
                "min_ats_score": 60,
                "auto_reject_below": None,
                "require_employee_review": True,
                "ats_enabled": True,
                "ai_verdict_enabled": True,
            }
    except Exception as e:
        logger.error(f"Error fetching AI config: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch AI config: {str(e)}")


@router.put("/config")
async def update_ai_config(
    request: AIConfigUpdate,
    session: AsyncSession = Depends(get_db),
    user=Depends(get_current_user),
):
    """
    Update company AI configuration settings.
    Only HR_MANAGER and SYSTEM_ADMIN can update.
    """
    try:
        from app.models.company_ai_config import CompanyAIConfig
        import uuid
        
        # Check permission (only HR or Admin can update)
        if user.role not in ["HR_MANAGER", "HR", "SYSTEM_ADMIN"]:
            raise HTTPException(status_code=403, detail="Only HR Manager or Admin can update AI settings")
        
        result = await session.execute(
            select(CompanyAIConfig).filter(CompanyAIConfig.company_id == user.company_id)
        )
        config = result.scalar_one_or_none()
        
        if not config:
            # Create new config
            config = CompanyAIConfig(
                id=uuid.uuid4(),
                company_id=user.company_id,
                min_passing_score=request.min_passing_score if request.min_passing_score is not None else 70,
                min_ats_score=request.min_ats_score if request.min_ats_score is not None else 60,
                auto_reject_below=request.auto_reject_below,
                require_employee_review=request.require_employee_review if request.require_employee_review is not None else True,
                ats_enabled=request.ats_enabled if request.ats_enabled is not None else True,
                ai_verdict_enabled=request.ai_verdict_enabled if request.ai_verdict_enabled is not None else True,
            )
            session.add(config)
        else:
            # Update existing config
            if request.min_passing_score is not None:
                config.min_passing_score = request.min_passing_score
            if request.min_ats_score is not None:
                config.min_ats_score = request.min_ats_score
            if request.auto_reject_below is not None:
                config.auto_reject_below = request.auto_reject_below
            if request.require_employee_review is not None:
                config.require_employee_review = request.require_employee_review
            if request.ats_enabled is not None:
                config.ats_enabled = request.ats_enabled
            if request.ai_verdict_enabled is not None:
                config.ai_verdict_enabled = request.ai_verdict_enabled
        
        await session.commit()
        await session.refresh(config)
        
        return config.to_dict()
    except HTTPException:
        raise
    except Exception as e:
        await session.rollback()
        logger.error(f"Error updating AI config: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to update AI config: {str(e)}")

